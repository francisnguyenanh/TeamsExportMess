"""
export_docx.py
==============
Tạo file .docx từ danh sách tin nhắn Teams.

Mỗi tin nhắn hiển thị:
  - Sender (bold, heading style)
  - Timestamp
  - Content (text)
  - Inline images (chèn ảnh trực tiếp)
  - File attachments (hiển thị link)
"""

import io
import re
import html as html_mod
import json
import tempfile
import requests
from pathlib import Path
from datetime import datetime, timezone, timedelta
from urllib.parse import urlparse

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── Constants ─────────────────────────────────────────────────────────────────

VN_TZ = timezone(timedelta(hours=7))
IMG_MAX_WIDTH = Inches(5.5)   # Chiều rộng tối đa ảnh trong doc
IMG_DOWNLOAD_TIMEOUT = 30     # Timeout download ảnh (s)

# Image extensions thường gặp
IMG_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tiff"}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _parse_dt_vn(iso_str: str) -> str:
    """Chuyển ISO datetime → chuỗi UTC+7."""
    if not iso_str:
        return ""
    try:
        dt = datetime.fromisoformat(iso_str.replace("Z", "+00:00"))
        dt_vn = dt.astimezone(VN_TZ)
        return dt_vn.strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        return iso_str


def _strip_html_keep_structure(html_str: str) -> list:
    """
    Parse HTML content → list of segments.
    Mỗi segment là dict:
      {"type": "text",  "value": "plain text"}
      {"type": "image", "url": "https://..."}
      {"type": "link",  "url": "https://...", "text": "display text"}
    """
    if not html_str:
        return [{"type": "text", "value": ""}]

    segments = []
    remaining = html_str

    # Pattern: tìm <img ...> và <a ...>...</a>
    tag_pattern = re.compile(
        r'(<img\s[^>]*?>)|(<a\s[^>]*?>(.*?)</a>)',
        re.IGNORECASE | re.DOTALL,
    )

    last_end = 0
    for m in tag_pattern.finditer(remaining):
        # Text trước tag
        before = remaining[last_end:m.start()]
        if before.strip():
            text = _clean_text(before)
            if text:
                segments.append({"type": "text", "value": text})

        if m.group(1):
            # <img> tag
            img_tag = m.group(1)
            src = _extract_attr(img_tag, "src")
            if src:
                segments.append({"type": "image", "url": src})
        elif m.group(2):
            # <a> tag
            a_tag = m.group(2)
            href = _extract_attr(a_tag, "href")
            link_text = _clean_text(m.group(3) or "")
            if href:
                segments.append({
                    "type": "link",
                    "url":  href,
                    "text": link_text or href,
                })

        last_end = m.end()

    # Text còn lại
    after = remaining[last_end:]
    if after.strip():
        text = _clean_text(after)
        if text:
            segments.append({"type": "text", "value": text})

    if not segments:
        segments.append({"type": "text", "value": _clean_text(html_str)})

    return segments


def _extract_attr(tag: str, attr_name: str) -> str:
    """Extract attribute value from an HTML tag string."""
    pattern = re.compile(
        rf'{attr_name}\s*=\s*["\']([^"\']*)["\']',
        re.IGNORECASE,
    )
    m = pattern.search(tag)
    return html_mod.unescape(m.group(1)) if m else ""


def _clean_text(html_fragment: str) -> str:
    """Loại bỏ HTML tags, decode entities, normalize whitespace."""
    # Block-level tags → newline
    text = re.sub(r"<br\s*/?>", "\n", html_fragment, flags=re.IGNORECASE)
    text = re.sub(r"</p>\s*", "\n", text, flags=re.IGNORECASE)
    text = re.sub(r"</div>\s*", "\n", text, flags=re.IGNORECASE)
    # Strip remaining tags
    text = re.sub(r"<[^>]+>", "", text)
    # Decode entities
    text = html_mod.unescape(text)
    # Normalize spaces
    text = text.replace("\u00a0", " ").replace("\u200b", "").replace("\ufeff", "")
    # Collapse blank lines
    lines = [line.strip() for line in text.splitlines()]
    cleaned = []
    blank = 0
    for line in lines:
        if line == "":
            blank += 1
            if blank <= 1:
                cleaned.append(line)
        else:
            blank = 0
            cleaned.append(line)
    return "\n".join(cleaned).strip()


def _download_image(url: str, headers: dict | None = None,
                    timeout: int = IMG_DOWNLOAD_TIMEOUT) -> bytes | None:
    """Download image từ URL, trả về bytes hoặc None nếu fail."""
    try:
        h = dict(headers) if headers else {}
        resp = requests.get(url, headers=h, timeout=timeout, stream=True)
        if resp.status_code != 200:
            return None
        content_type = resp.headers.get("Content-Type", "")
        if "image" not in content_type and "octet-stream" not in content_type:
            # Check by content (first bytes)
            data = resp.content
            if data[:4] in (b"\x89PNG", b"\xff\xd8\xff", b"GIF8", b"RIFF", b"BM"):
                return data
            return None
        return resp.content
    except Exception:
        return None


def _is_image_url(url: str) -> bool:
    """Check if URL likely points to an image."""
    parsed = urlparse(url)
    path_lower = parsed.path.lower()
    # Direct file extension check
    if any(path_lower.endswith(ext) for ext in IMG_EXTS):
        return True
    # Teams/SharePoint image hosting patterns
    if "getpersonalizedblob" in path_lower or "imgo" in path_lower:
        return True
    # Azure Media Services
    if "ams/" in path_lower or "objects/" in path_lower:
        return True
    return False


def _is_onedrive_or_sharepoint_url(url: str) -> bool:
    """Check if URL is a OneDrive/SharePoint file link."""
    domain = urlparse(url).netloc.lower()
    return any(d in domain for d in [
        "sharepoint.com", "onedrive.com", "1drv.ms",
        "sharepoint-df.com",
    ])


# ── Main writer ───────────────────────────────────────────────────────────────

def parse_chatsvc_message(msg: dict) -> dict | None:
    """
    Parse a raw chatsvc message → dict for docx export.
    Returns None for system/skip messages.

    Output format:
    {
        "sender":      "Display Name",
        "datetime":    "2026-03-27 10:30:00",
        "content_raw": "<html content>",
        "segments":    [{"type": "text"|"image"|"link", ...}],
        "attachments": [{"name": "file.xlsx", "url": "https://..."}],
        "images":      ["https://img1", ...],
        "message_id":  "xxx",
    }
    """
    msg_type = msg.get("messagetype", msg.get("messageType", ""))

    # Skip system events
    if msg_type in ("ThreadActivity/AddMember", "ThreadActivity/DeleteMember",
                    "ThreadActivity/TopicUpdate", "ThreadActivity/MemberJoined",
                    "ThreadActivity/MemberLeft"):
        return None

    # Accept known message types, or anything with content
    if msg_type not in ("Text", "RichText/Html", "RichText", "text", "message"):
        if not msg.get("content") and not msg.get("body"):
            return None

    # ── Sender ────────────────────────────────────────────────────────────
    sender = "Unknown"
    imd = msg.get("imdisplayname", "")
    if imd:
        sender = imd
    else:
        from_mri = msg.get("from", "")
        if isinstance(from_mri, str):
            sender = from_mri.split(":")[-1][:30]
        elif isinstance(from_mri, dict):
            user = from_mri.get("user") or from_mri.get("application") or {}
            sender = user.get("displayName", "Unknown")

    # ── Timestamp ─────────────────────────────────────────────────────────
    composed = (msg.get("composetime")
                or msg.get("createdDateTime")
                or msg.get("originalarrivaltime", ""))
    dt_str = _parse_dt_vn(composed)

    # ── Content (raw HTML) ────────────────────────────────────────────────
    content_raw = msg.get("content", "")
    if not content_raw:
        body = msg.get("body", {})
        if isinstance(body, dict):
            content_raw = body.get("content", "")
        elif isinstance(body, str):
            content_raw = body

    # Skip system event markers
    if not content_raw or content_raw.strip() in ("<systemEventMessage/>", ""):
        return None

    # ── Parse content → segments ──────────────────────────────────────────
    segments = _strip_html_keep_structure(content_raw)

    # ── Extract images from various sources ───────────────────────────────
    images = []

    # 1. Images from <img src="..."> in content
    img_pattern = re.compile(r'<img\s[^>]*?src=["\']([^"\']+)["\']', re.IGNORECASE)
    for m in img_pattern.finditer(content_raw):
        img_url = html_mod.unescape(m.group(1))
        if img_url and img_url not in images:
            images.append(img_url)

    # 2. amsreferences — Azure Media Service (inline images/screenshots)
    ams_refs = msg.get("amsreferences", [])
    if isinstance(ams_refs, list):
        for ref in ams_refs:
            if isinstance(ref, str) and ref not in images:
                images.append(ref)

    # 3. properties.files (chatsvc specific)
    props = msg.get("properties", {})
    if isinstance(props, dict):
        files_str = props.get("files", "")
        if files_str:
            try:
                if isinstance(files_str, str):
                    files_list = json.loads(files_str)
                else:
                    files_list = files_str
                if isinstance(files_list, list):
                    for f in files_list:
                        if isinstance(f, dict):
                            furl = f.get("objectUrl") or f.get("url", "")
                            if furl and _is_image_url(furl) and furl not in images:
                                images.append(furl)
            except (json.JSONDecodeError, TypeError):
                pass

    # ── Extract attachments (files / OneDrive links) ──────────────────────
    attachments = []

    # From <a> tags that link to SharePoint/OneDrive
    link_pattern = re.compile(
        r'<a\s[^>]*?href=["\']([^"\']+)["\'][^>]*?>(.*?)</a>',
        re.IGNORECASE | re.DOTALL,
    )
    for m in link_pattern.finditer(content_raw):
        href = html_mod.unescape(m.group(1))
        link_text = _clean_text(m.group(2))
        if _is_onedrive_or_sharepoint_url(href):
            attachments.append({"name": link_text or "File", "url": href})

    # From properties.files (non-image files)
    if isinstance(props, dict):
        files_str = props.get("files", "")
        if files_str:
            try:
                if isinstance(files_str, str):
                    files_list = json.loads(files_str)
                else:
                    files_list = files_str
                if isinstance(files_list, list):
                    for f in files_list:
                        if isinstance(f, dict):
                            furl = f.get("objectUrl") or f.get("url", "")
                            fname = f.get("fileName") or f.get("title", "File")
                            if furl and not _is_image_url(furl):
                                attachments.append({"name": fname, "url": furl})
            except (json.JSONDecodeError, TypeError):
                pass

    msg_id = msg.get("id", msg.get("messageId", msg.get("clientmessageid", "")))

    return {
        "sender":      sender,
        "datetime":    dt_str,
        "content_raw": content_raw,
        "segments":    segments,
        "attachments": attachments,
        "images":      images,
        "message_id":  str(msg_id),
    }


def write_docx(messages: list[dict], output_path: Path, chat_name: str,
               auth_headers: dict | None = None,
               log_fn=None) -> int:
    """
    Tạo file .docx từ danh sách tin nhắn đã parse.

    Args:
        messages:     List of parsed message dicts (from parse_chatsvc_message)
        output_path:  Path to output .docx file
        chat_name:    Tên group chat (hiển thị ở tiêu đề)
        auth_headers: Headers để download ảnh (Bearer token)
        log_fn:       Optional log callback

    Returns:
        Số tin nhắn đã viết
    """
    def log(msg):
        if log_fn:
            log_fn(msg)

    doc = Document()

    # ── Style setup ───────────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # ── Title ─────────────────────────────────────────────────────────────
    title = doc.add_heading(chat_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Export timestamp
    now_vn = datetime.now(VN_TZ).strftime("%Y-%m-%d %H:%M:%S UTC+7")
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"Exported: {now_vn}")
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    count = 0
    img_cache = {}  # url → bytes (avoid re-downloading)

    for msg in messages:
        sender = msg.get("sender", "Unknown")
        dt_str = msg.get("datetime", "")
        segments = msg.get("segments", [])
        images = msg.get("images", [])
        attachments = msg.get("attachments", [])

        # ── Sender (bold heading) ─────────────────────────────────────────
        sender_para = doc.add_heading(sender, level=3)

        # ── Timestamp (gray, smaller font) ────────────────────────────────
        ts_para = doc.add_paragraph()
        ts_run = ts_para.add_run(dt_str)
        ts_run.font.size = Pt(8)
        ts_run.font.color.rgb = RGBColor(128, 128, 128)
        ts_run.font.italic = True

        # ── Content segments ──────────────────────────────────────────────
        content_para = doc.add_paragraph()
        has_content = False

        for seg in segments:
            seg_type = seg.get("type", "text")

            if seg_type == "text":
                text = seg.get("value", "")
                if text:
                    content_para.add_run(text)
                    has_content = True

            elif seg_type == "link":
                url = seg.get("url", "")
                text = seg.get("text", url)
                if url:
                    # Add as clickable hyperlink text
                    run = content_para.add_run(f"{text}")
                    run.font.color.rgb = RGBColor(0, 102, 204)
                    run.font.underline = True
                    # Add URL in parentheses if different from display text
                    if text != url and not text.startswith("http"):
                        url_run = content_para.add_run(f" ({url})")
                        url_run.font.size = Pt(7)
                        url_run.font.color.rgb = RGBColor(150, 150, 150)
                    has_content = True

            elif seg_type == "image":
                # Will be handled in the images section below
                pass

        if not has_content:
            plain = _clean_text(msg.get("content_raw", ""))
            if plain:
                content_para.add_run(plain)

        # ── Inline Images ─────────────────────────────────────────────────
        img_inserted = 0
        for img_url in images:
            if not img_url:
                continue

            # Try to download and embed
            img_data = None
            if img_url in img_cache:
                img_data = img_cache[img_url]
            else:
                # Build headers for image download
                dl_headers = {}
                if auth_headers:
                    dl_headers.update(auth_headers)

                # AMS images need the chatsvc token
                img_data = _download_image(img_url, headers=dl_headers)
                img_cache[img_url] = img_data

            if img_data:
                try:
                    img_stream = io.BytesIO(img_data)
                    img_para = doc.add_paragraph()
                    img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = img_para.add_run()
                    run.add_picture(img_stream, width=IMG_MAX_WIDTH)
                    img_inserted += 1
                except Exception as e:
                    log(f"⚠️ Không thể chèn ảnh: {e}")
                    # Fallback: add link
                    p = doc.add_paragraph()
                    r = p.add_run(f"📷 [Image] {img_url}")
                    r.font.color.rgb = RGBColor(0, 102, 204)
                    r.font.size = Pt(8)
            else:
                # Cannot download → show link
                p = doc.add_paragraph()
                r = p.add_run(f"📷 [Image — download failed] {img_url}")
                r.font.color.rgb = RGBColor(200, 100, 100)
                r.font.size = Pt(8)

        # ── File Attachments (links) ──────────────────────────────────────
        for att in attachments:
            att_name = att.get("name", "File")
            att_url = att.get("url", "")
            p = doc.add_paragraph()
            r = p.add_run(f"📎 {att_name}")
            r.font.bold = True
            r.font.size = Pt(9)
            if att_url:
                url_run = p.add_run(f"\n   {att_url}")
                url_run.font.color.rgb = RGBColor(0, 102, 204)
                url_run.font.size = Pt(8)
                url_run.font.underline = True

        # ── Separator ─────────────────────────────────────────────────────
        sep = doc.add_paragraph()
        sep_run = sep.add_run("─" * 60)
        sep_run.font.size = Pt(6)
        sep_run.font.color.rgb = RGBColor(200, 200, 200)

        count += 1

    # ── Save ──────────────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log(f"📄 Saved {count} messages → {output_path.name}")

    return count
